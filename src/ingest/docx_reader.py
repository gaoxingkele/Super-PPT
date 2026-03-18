# -*- coding: utf-8 -*-
"""DOCX 文件读取：提取文本、表格、内嵌图片。"""
from pathlib import Path


def read_docx(file_path: Path, image_dir: Path = None) -> dict:
    """
    读取 DOCX 文件，返回结构化内容。

    Args:
        file_path: DOCX 文件路径
        image_dir: 图片保存目录（None 则不提取图片）

    Returns:
        {
            "text": str,
            "meta": dict,          # {title}
            "tables": list[dict],  # [{headers, rows}, ...]
            "images": list[str],   # 保存的图片路径列表
        }
    """
    from docx import Document

    doc = Document(str(file_path))

    # 提取文本
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # 保留标题层级
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    level = int(level)
                except ValueError:
                    level = 1
                paragraphs.append(f"{'#' * level} {text}")
            else:
                paragraphs.append(text)

    full_text = "\n\n".join(paragraphs)

    # 提取表格
    tables = []
    for table in doc.tables:
        rows_data = []
        for row in table.rows:
            rows_data.append([cell.text.strip() for cell in row.cells])
        if rows_data:
            tables.append({"headers": rows_data[0], "rows": rows_data[1:]})

    # 提取内嵌图片
    images = []
    if image_dir:
        image_dir = Path(image_dir)
        image_dir.mkdir(parents=True, exist_ok=True)
        import zipfile
        img_counter = 0
        with zipfile.ZipFile(str(file_path), "r") as z:
            for name in z.namelist():
                if name.startswith("word/media/"):
                    img_counter += 1
                    ext = Path(name).suffix or ".png"
                    out_name = f"img_{img_counter:03d}{ext}"
                    out_path = image_dir / out_name
                    out_path.write_bytes(z.read(name))
                    images.append(str(out_path))

    # 元数据
    meta = {}
    core = doc.core_properties
    if core.title:
        meta["title"] = core.title

    return {"text": full_text, "meta": meta, "tables": tables, "images": images}
